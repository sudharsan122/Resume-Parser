#without LOGO
# app.py  (your Streamlit app with integrated Gemini skill+experience extraction)
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import io
import re
import docx
import pdfplumber
import zipfile
import streamlit as st
import pandas as pd
import time
import json
from datetime import datetime

# ------------------ HARD-CODED GEMINI KEY (replace locally) ------------------

GEMINI_API_KEY = "AIzaSyAVf_3yIxaO9ecFyWp7SIdbmgf4YnXGer8"

# ------------------ genai client ------------------
try:
    from google import genai
except Exception:
    genai = None

if genai is None:
    raise RuntimeError("Install google-genai: pip install google-genai")

client = genai.Client(api_key=GEMINI_API_KEY)

# ============================================================== #
# GLOBAL DEBUG REGISTRY (for clean, structured logs)
# ============================================================== #
EXTRACT_DEBUG_REGISTRY = {}  # { filename: {...} }

# ------------------ LLM helpers (skills + years) ------------------

MAX_CHARS = 14000

def call_gemini_for_skills(resume_text: str, model: str = "gemini-2.5-flash", max_retries: int = 2):
    """
    Returns (skills_list, raw_output).
    skills_list: list[str] lowercased tokens (may be empty).
    """
    if not resume_text:
        return [], "<empty>"

    txt = resume_text if len(resume_text) <= MAX_CHARS else resume_text[:MAX_CHARS]
    prompt = f"""
You are an extractor. Given the resume text below, return ONLY a single JSON object:

{{"skills": ["skill1","skill2", ...]}}

Rules:
- Return short canonical skill tokens like "python", "c++", "embedded linux", "device tree", "u-boot", "yocto", "i2c", "spi", "git".
- Normalize common variants (react.js -> react, node js -> node.js, powerbi -> power bi).
- Deduplicate and include only skills actually present in the resume.
- Do NOT include company names, addresses, or long descriptive sentences.
- Output EXACTLY one JSON object and nothing else.

Resume:
\"\"\"{txt}\"\"\"
""".strip()

    for attempt in range(max_retries + 1):
        try:
            resp = client.models.generate_content(model=model, contents=prompt)
            raw = resp.text if hasattr(resp, "text") else str(resp)
            # find JSON object
            m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
            if m:
                try:
                    obj = json.loads(m.group(0))
                    skills = obj.get("skills", [])
                    if isinstance(skills, list):
                        # normalize & dedupe preserving order
                        cleaned = []
                        seen = set()
                        for s in skills:
                            if not isinstance(s, str):
                                continue
                            tok = s.strip().lower()
                            tok = tok.replace('react.js', 'react').replace('reactjs', 'react')
                            tok = tok.replace('node js', 'node.js').replace('nodejs', 'node.js')
                            tok = tok.replace('powerbi', 'power bi')
                            tok = re.sub(r'[\._]+', ' ', tok)
                            tok = re.sub(r'\s+', ' ', tok).strip()
                            if tok and tok not in seen:
                                seen.add(tok)
                                cleaned.append(tok)
                        return cleaned, raw
                except Exception:
                    pass

            # fallback: try to parse an array literal
            arr_match = re.search(r'\[\s*([^\]]+?)\s*\]', raw, flags=re.DOTALL)
            if arr_match:
                items = re.findall(r'["\']([^"\']+)["\']', arr_match.group(0))
                cleaned = []
                seen = set()
                for s in items:
                    tok = s.strip().lower()
                    tok = tok.replace('react.js', 'react').replace('reactjs', 'react')
                    tok = tok.replace('node js', 'node.js').replace('nodejs', 'node.js')
                    tok = re.sub(r'[\._]+', ' ', tok)
                    tok = re.sub(r'\s+', ' ', tok).strip()
                    if tok and tok not in seen:
                        seen.add(tok)
                        cleaned.append(tok)
                if cleaned:
                    return cleaned, raw

            # If we didn't get JSON, fallback to empty so caller will use local fallback
            return [], raw
        except Exception as e:
            # transient retry
            if attempt < max_retries:
                time.sleep(0.6 * (attempt + 1))
                continue
            return [], f"<error: {e}>"

def call_gemini_for_years(resume_text: str, model: str = "gemini-2.5-flash", max_retries: int = 2):
    """
    Returns (float_years, raw_output)
    Expects LLM to return ONLY {"total_years": <float>}
    """
    if not resume_text:
        return 0.0, "<empty>"

    txt = resume_text if len(resume_text) <= MAX_CHARS else resume_text[:MAX_CHARS]
    today = datetime.now().date()
    prompt = f"""
You are a strict extractor. Given the resume text below, compute the candidate's total professional work experience in years, MERGING overlapping jobs so they are not double-counted. Return ONLY a single JSON object with this numeric field:

{{"total_years": <float>}}

Rules:
- Count full months; express years as decimal with one digit after decimal (e.g., 3.5).
- Treat "present/current" as up to today's date ({today}).
- Merge overlapping jobs before summing so overlaps are not double-counted.
- If you cannot detect any valid dates or experience, return {{"total_years": 0.0}}.
- DO NOT output any text except the single JSON object above.

Resume:
\"\"\"{txt}\"\"\" 
""".strip()

    for attempt in range(max_retries + 1):
        try:
            resp = client.models.generate_content(model=model, contents=prompt)
            raw = resp.text if hasattr(resp, "text") else str(resp)
            m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
            if m:
                try:
                    obj = json.loads(m.group(0))
                    value = float(obj.get("total_years", 0.0))
                    return round(value, 1), raw
                except Exception:
                    pass
            # try to find numeric fallback
            m2 = re.search(r'(\d+(?:\.\d+)?)', raw)
            if m2:
                try:
                    return round(float(m2.group(1)), 1), raw
                except:
                    pass
            return 0.0, raw
        except Exception as e:
            if attempt < max_retries:
                time.sleep(0.6 * (attempt + 1))
                continue
            return 0.0, f"<error: {e}>"

# ------------------ TEXT extraction utilities (existing) ------------------
def extract_text_from_pdf(pdf_file):
    try:
        pdf_file.seek(0)
        with pdfplumber.open(pdf_file) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages)
            text = re.sub(r'\s+', ' ', text).strip()
            text = re.sub(r'\(https?:\/\/\S+\)', '', text)
            return text.lower()
    except Exception:
        return ""

def extract_text_from_docx(docx_file):
    """Paragraph + table text."""
    try:
        docx_file.seek(0)
        document = docx.Document(docx_file)
        parts = []
        for p in document.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        parts.append(ct)
        text = "\n".join(parts)
        text = re.sub(r'\s+', ' ', text).strip()
        return text.lower()
    except Exception:
        return ""

def extract_text_from_txt(txt_file):
    try:
        txt_file.seek(0)
        raw = txt_file.getvalue()
        if isinstance(raw, bytes):
            raw = raw.decode('utf-8', errors='ignore')
        text = str(raw)
        text = re.sub(r'\s+', ' ', text).strip()
        return text.lower()
    except Exception:
        return ""

# ------------------ Hardened Skill extraction (drop-in replacement, with technical-signal guard) ------------------

BASE_KEYWORDS = [
    "c", "c++", "c#", "python", "java", "javascript", "typescript", "go", "rust", "ruby", "php", "scala", "kotlin", "swift", "r",
    "react", "angular", "vue", "next.js", "svelte", "html", "css", "sass", "tailwind",
    "node.js", "express", "django", "flask", "spring boot", "spring", "laravel", "asp.net",
    "sql", "postgresql", "mysql", "mongodb", "redis", "oracle", "mssql", "cassandra",
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible", "helm",
    "jenkins", "github actions", "gitlab-ci", "circleci",
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "xgboost", "lightgbm", "nlp", "opencv", "spacy",
    "spark", "hadoop", "etl", "airflow",
    "embedded linux", "yocto", "petalinux", "u-boot", "device tree", "kernel", "linux kernel", "bsp",
    "arm", "raspberry pi", "stm32", "nxp", "imx", "qualcomm",
    "i2c", "spi", "uart", "gpio", "pcie", "usb", "ethernet", "can", "i2s",
    "board bring-up", "firmware", "bootloader", "driver development", "kernel drivers", "device drivers",
    "git", "gdb", "cmake", "make", "gcc", "clang", "vivado", "quartus", "jtag",
    "linux", "bash", "shell", "systemd", "excel", "tableau", "power bi", "docker-compose",
    "openwrt", "buildroot", "qemu", "tpm", "jtag"
]

NORMALIZE_MAP = {
    "react.js": "react",
    "reactjs": "react",
    "nodejs": "node.js",
    "node js": "node.js",
    "powerbi": "power bi",
    "u boot": "u-boot",
    "u_boot": "u-boot",
    "device-tree": "device tree",
    "embedded c": "c",
    "c plus plus": "c++",
    "cplusplus": "c++",
    "usb 3 0": "usb 3.0",
    "usb3.0": "usb 3.0",
    "wi fi": "wi-fi",
    "i 2 c": "i2c",
    "i 2 s": "i2s",
    "yocto project": "yocto",
    "petalinux sdk": "petalinux",
    "system verilog": "systemverilog",
    "devops": "devops",
    "open source": "open source",
    "open-source": "open source",
    "kernel internals": "linux kernel internals",
    "build systems": "build systems",
    "boot time": "boot time optimization",
    "firmware update": "firmware updates",
    "firmware updates": "firmware updates",
    "linux kernel internals -": "linux kernel internals"
}

# Aggressive stop words / tokens
STOP = {
    'a','an','the','and','or','of','for','to','in','on','with','without','as','by','at','from','into','over',
    'is','are','was','were','be','being','been','can','could','should','would','may','might','must',
    'that','this','these','those','it','its','their','our','your','you','we','they',
    'years','year','experience','exp','minimum','at','least','plus','including','etc','such','as',
    'hands','on','hands-on','proficient','proficiency','expertise','expert','knowledge','understanding',
    'ability','capability','skills','skill','good','strong','excellent','great',
    'preferred','nice','have','nice-to-have','must','required','requirement','requirements',
    'responsibilities','responsibility','role','roles','about','job','description','candidate',
    'team','teams','collaboration','communicate','communication','problem','solving','ownership',
    'drive','driven','self','motivated','fast','paced','fast-paced','environment',
    'worked','working','work','design','develop','development','implement','implementation',
    'architecture','architectures','across','using','use','used','involves','contribute','contributions',
    'review','reviews','code','coding','test','testing','debug','debugging','issue','issues','resolve','resolution',
    'preferred qualifications', 'qualifications', 'key', 'preferred', 'required', 'related field', 'related fields',
    's or master', 's or masters', 's or master s'
}

VERB_LIKE = {'collaborate','integrate','lead','support','join','ensure','work','maintain','validate','resolve','contribute','contributes','contributing','seeking','responsible','responsibilities','will','must','should','prefer','supporting'}

# Short list of technical indicator substrings (keeps tokens that mention these)
TECH_INDICATORS = {
    'linux','kernel','yocto','u-boot','openwrt','buildroot','docker','qemu','jtag','gdb',
    'spi','i2c','uart','gpio','ethernet','usb','wi-fi','wifi','lte','tpm','bsp','boot',
    'bootloader','firmware','driver','broadcom','mediatek','qualcomm','nxp','arm','imx',
    'stm32','raspberry','gcc','cmake','make','gcc','python','java','c++','c#','node.js',
    'react','spark','hadoop','pandas','numpy','tensorflow','pytorch','redis','postgresql',
    'mysql','mongodb','aws','azure','gcp','kubernetes','terraform','ansible','helm','jenkins'
}

def dedupe_preserve_order(items):
    seen = set()
    out = []
    for x in items:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out

def normalize_token(tok: str) -> str:
    if not tok:
        return tok
    s = tok.strip().lower()
    # remove extraneous punctuation around words
    s = re.sub(r'^[\-\‚Ä¢\*\‚Ä¢\s]+', '', s)
    s = re.sub(r'[\_\t]+', ' ', s)
    s = re.sub(r'[\u2010-\u2015]', '-', s)
    s = s.replace('&', ' and ')
    s = re.sub(r'[\(\)\[\]]', ' ', s)
    s = re.sub(r'\s*[\/\\]+\s*', ' / ', s)
    s = s.replace(',', ' ')
    s = re.sub(r'\s+', ' ', s).strip()
    # apply mapping conservatively (whole-word)
    for k, v in NORMALIZE_MAP.items():
        if s == k or re.search(r'\b' + re.escape(k) + r'\b', s):
            s = re.sub(r'\b' + re.escape(k) + r'\b', v, s)
    # remove stray standalone punctuation
    s = re.sub(r'(?<!\d)\.(?!\d)', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    # final small heuristics
    if s == 'node':
        s = 'node.js'
    return s

def _preclean_text(text: str) -> str:
    """Remove markdown headers, bullets, long dash lines and other template noise."""
    if not text:
        return ""
    t = text
    # drop markdown headings and separators
    t = re.sub(r'(?m)^\s*#{1,6}\s*', ' ', t)        # remove leading ### etc.
    t = re.sub(r'(?m)^\s*[-‚Äì‚Äî]{3,}\s*$', ' ', t)    # lines with --- or similar
    # remove common bullet markers at line starts
    t = re.sub(r'(?m)^\s*[\-\*\‚Ä¢]\s*', ' ', t)
    # remove "### preferred qualifications" type full-line markers
    t = re.sub(r'(?mi)\b(preferred qualifications|preferred|key qualifications|key|qualifications)\b', ' ', t)
    # remove common "e.g." variants to make splitting easier
    t = re.sub(r'\be\.g\.\b', ' e g ', t)
    t = re.sub(r'\bE\.g\.\b', ' e g ', t)
    # trim repeated separators
    t = re.sub(r'\s+', ' ', t)
    return t.strip()

def _split_candidate_into_parts(candidate: str) -> list:
    """
    When a candidate token contains multiple items separated by slashes, commas, 'e g', ' and ', '-' etc,
    split into smaller pieces and normalize each piece.
    """
    if not candidate:
        return []
    # split on common separators that join multiple skills
    parts = re.split(r'[,/]| e g | and | & | \| |;|-{2,}|\s-\s|\/', candidate)
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # further split tokens like "usb ethernet" where both are skills
        # but only if both are short tokens (heuristic)
        words = p.split()
        if len(words) > 1 and len(words) <= 3:
            # if all words look like separate known base keywords, split them
            subparts = []
            for w in words:
                w_norm = normalize_token(w)
                subparts.append(w_norm)
            # if at least one subpart looks good, return them
            out.extend([sp for sp in subparts if sp])
        else:
            out.append(p)
    return out

def token_is_noisy(token: str, EXTRA_STOP_CONTAINS=None) -> bool:
    if EXTRA_STOP_CONTAINS is None:
        EXTRA_STOP_CONTAINS = {
            'job', 'description', 'role', 'position', 'candidate', 'responsibilities', 'requirements',
            'years', 'experience', '3+', '5+', 'month', 'months', 'join', 'we', 'will', 'ensure',
            'work', 'working', 'seeking', 'apply', 'must', 'required', 'prefer', 'good', 'strong', 'ability',
            'responsibilities', 'responsibility', 'will be', 'we are', 'job description', 'preferred qualifications',
            'related field', 'qualifications', 'new', 'for new'
        }
    t = token.strip().lower()
    if not t:
        return True
    if len(t) <= 1:
        return True
    if re.fullmatch(r'[\d\+\-\.]+', t):
        return True
    # drop if contains boilerplate / template fragments
    for bad in EXTRA_STOP_CONTAINS:
        if f' {bad} ' in f' {t} ' or t.startswith(bad + ' ') or t.endswith(' ' + bad):
            return True
    # drop tokens that are mostly stop words
    words_t = re.split(r'[\s/\\]+', t)
    alpha_words = [w for w in words_t if re.search(r'[a-z]', w)]
    if len(alpha_words) == 0:
        return True
    stop_count = sum(1 for w in alpha_words if w in STOP)
    if stop_count / max(1, len(alpha_words)) > 0.5:
        return True
    # drop verb-like tokens
    if t in VERB_LIKE:
        return True
    # length guard
    if len(t) > 60:
        return True
    return False

# ------------------ Patch: expand technical signals to preserve multiword skills ------------------

# add this near the top with other constants
EXTRA_BASE_PHRASES = {
    "board bring-up",
    "boot time optimization",
    "build systems",
    "code reviews",
    "containerization",
    "continuous integration",
    "debugging tools",
    "device tree",
    "documentation",
    "embedded linux",
    "embedded systems",
    "firmware updates",
    "init scripts",
    "kernel drivers",
    "linux kernel internals",
    "logic analyzers",
    "memory optimization",
    "networking",
    "open source",
    "oscilloscope",
    "router platforms",
    "secure boot",
    "shell scripting",
    "system performance optimization",
    "test automation",
    "virtualization",
    "wireless networking",
    "production-ready",
    "build systems",
    "package platform",
}

# extend TECH_INDICATORS with a few more substrings (if you already have TECH_INDICATORS, update it)
TECH_INDICATORS.update({
    'optimization','optimization','performance','automation','scripting','integration','integration pipelines',
    'containerization','continuous','ci','cd','debug','debugging','analysis','analytics','documentation',
    'network','networking','platform','platforms','memory','performance','test','testing','automation','secure','security'
})

# suffixes or keywords that when present indicate a technical phrase (helps multiword detection)
TECH_SUFFIXES = {
    'optimization','automation','systems','scripting','integration','containerization','virtualization',
    'performance','debugging','testing','analysis','documentation','networking','security','drivers','bring-up',
    'bringup','updates','update','platform','platforms','tools','engineering','engineering','devops'
}

def token_is_technical(token: str) -> bool:
    """
    Smarter technical check:
      - exact match in BASE_KEYWORDS OR EXTRA_BASE_PHRASES OR
      - contains a TECH_INDICATOR OR
      - any word contains an indicator OR suffix (like 'optimization', 'scripting')
      - or token contains tech punctuation like + # . (c++, c#)
    """
    t = token.strip().lower()
    if not t:
        return False

    # direct whitelist (single word or canonical multi-word)
    if t in BASE_KEYWORDS or t in EXTRA_BASE_PHRASES:
        return True

    # direct substring indicators
    for ind in TECH_INDICATORS:
        if ind in t:
            return True

    # suffix match on words (helpful for multi-word tokens)
    for w in re.split(r'[\s/\\]+', t):
        if w in BASE_KEYWORDS or w in EXTRA_BASE_PHRASES:
            return True
        for suf in TECH_SUFFIXES:
            if w.endswith(suf) or suf in w:
                return True
        for ind in TECH_INDICATORS:
            if ind in w:
                return True

    # keep tokens that include special tech characters and alpha numeric content
    if any(ch in t for ch in ['+','#','.']):
        if re.search(r'[a-z0-9]', t):
            return True

    return False

def get_keywords_from_jd(jd_text: str, max_terms: int = 150) -> list:
    if not jd_text:
        return []

    text = _preclean_text(jd_text.lower())

    # keep only relevant chars
    text = re.sub(r'(https?://\S+|www\.\S+)', ' ', text)
    text = re.sub(r'[^a-z0-9+\#\.\-/\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    # 1) BASE_KEYWORDS high-precision scan
    found = []
    for kw in BASE_KEYWORDS:
        patt = r'\b' + re.escape(kw) + r'\b'
        if re.search(patt, text):
            found.append(kw)

    # 2) candidate generation (1-3 grams)
    token_candidates = set()
    words = text.split()
    L = len(words)
    for i in range(L):
        token_candidates.add(words[i])
        for l in (2, 3):
            if i + l <= L:
                seq = " ".join(words[i:i + l])
                token_candidates.add(seq)

    # normalize candidates and split combined tokens
    normalized = []
    for tok in token_candidates:
        # split if tok contains multiple items (e.g., "bootloaders u-boot device")
        parts = _split_candidate_into_parts(tok)
        for p in parts:
            p_norm = normalize_token(p)
            if not p_norm:
                continue
            if len(p_norm) < 2:
                continue
            if p_norm.isdigit():
                continue
            if not re.search(r'[a-z]', p_norm):
                continue
            if p_norm in STOP:
                continue
            # remove leading connector words
            p_norm = re.sub(r'^(and|or|with|the|a|an)\s+', '', p_norm).strip()
            normalized.append(p_norm)

    # combine base hits and normalized candidates
    combined = found + normalized
    combined = [c.strip() for c in combined if c and len(c) <= 60]
    combined = dedupe_preserve_order(combined)

    # final stronger cleanup & technical guard
    final = []
    for s in combined:
        # split tricky combined phrases again and normalize pieces
        parts = _split_candidate_into_parts(s)
        for p in parts:
            p2 = normalize_token(p)
            if not p2:
                continue
            p2 = re.sub(r'\s+', ' ', p2).strip()
            p2 = re.sub(r'^(and|or|with|the|a|an)\s+', '', p2).strip()
            # drop noisy / boilerplate
            if token_is_noisy(p2):
                continue
            # require technical signal (either base keyword or technical indicator)
            if not token_is_technical(p2):
                continue
            if p2 in STOP:
                continue
            if p2 in VERB_LIKE:
                continue
            final.append(p2)

    final = dedupe_preserve_order(final)[:max_terms]
    return set(final)

def get_skills_from_text(text: str) -> list:
    """Wrapper that returns an ordered list of cleaned skills."""
    return get_keywords_from_jd(text, max_terms=500)

# ----------------------------------------------------------------------------------------------------------------------
# EXPERIENCE EXTRACTION (existing local code) - keep unchanged
# ----------------------------------------------------------------------------------------------------------------------
def parse_date_any(s):
    s = s.strip()
    s = re.sub(r'\s+', ' ', s)
    fmts = ("%b %Y", "%B %Y", "%Y/%m", "%m/%Y", "%Y")
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt)
        except:
            continue
    m = re.match(r'^\d{4}$', s)
    if m:
        try:
            return datetime.strptime(s, "%Y")
        except:
            pass
    return None

def extract_experience_from_resume(text, filename="", aggressive_edu=True):
    # (copy the full function from your existing code above)
    # For brevity here, we call the same function content you already included earlier.
    # Ensure you paste the same function body here (we've kept it in your file).
    # I'm including the original body for correctness (already in your file).
    fname = filename or "unknown"
    dbg = {
        "source": "none",
        "explicit_values": [],
        "structured_intervals": [],
        "date_range_intervals": [],
        "final_years": 0.0,
    }

    print("\n" + "‚ïê"*78)
    print(f"üßæ RESUME EXPERIENCE DEBUG ‚Äî {fname}")
    print("‚ïê"*78)

    if not text:
        print("‚ö†Ô∏è No text extracted from resume. Returning 0.0 years.")
        EXTRACT_DEBUG_REGISTRY[fname] = dbg
        return 0.0

    text_low = text.lower()

    explicit_values = []
    # 1) Explicit numeric phrases
    for m in re.finditer(r'\b(\d+(?:\.\d+)?)\s*(?:\+)?\s*(?:years?|yrs?|y)\b', text_low):
        try:
            val = float(m.group(1))
            explicit_values.append(val)
            print(f"üîé Found explicit years mention: {val}")
        except:
            pass

    for m in re.finditer(r'\b(\d+)\s*[yY]\s*(\d+)\s*[mM]\b', text_low):
        try:
            val = int(m.group(1)) + int(m.group(2))/12.0
            explicit_values.append(val)
            print(f"üîé Found Y-M pattern => {val:.2f} years")
        except:
            pass

    for m in re.finditer(r'(\d+(?:\.\d+)?)\s+years\s+of\s+experience', text_low):
        try:
            val = float(m.group(1))
            explicit_values.append(val)
            print(f"üîé Found 'years of experience' => {val}")
        except:
            pass

    if explicit_values:
        result = round(max(explicit_values), 1)
        dbg["source"] = "explicit_number"
        dbg["explicit_values"] = explicit_values[:]
        dbg["final_years"] = result
        print(f"‚úÖ Using explicit numeric mention ‚Üí {result} years (max of {explicit_values})")
        EXTRACT_DEBUG_REGISTRY[fname] = dbg
        return result

    # 2) Section-focused parsing (Experience vs Education)
    lines = [ln.strip() for ln in re.split(r'[\r\n]+', text_low) if ln.strip()]
    current_section = None
    experience_text_lines = []

    education_keywords = re.compile(r'\b(education|academic|university|college|degree|school|certification)\b', re.I)
    experience_keywords = re.compile(r'\b(experience|employment|work history|professional experience|work experience|career)\b', re.I)

    for idx, ln in enumerate(lines):
        if experience_keywords.search(ln):
            current_section = "experience"
            print(f"üìå Section ‚Üí experience (line {idx})")
            continue
        if education_keywords.search(ln):
            current_section = "education"
            print(f"üìå Section ‚Üí education (line {idx})")
            continue

        if ln.isupper() and len(ln.split()) <= 4:
            if 'experience' in ln.lower() or 'employment' in ln.lower() or 'work' in ln.lower():
                current_section = "experience"
                print(f"üìå Section ‚Üí EXPERIENCE (header, line {idx})")
                continue
            if 'education' in ln.lower() or 'academic' in ln.lower():
                current_section = "education"
                continue

        if aggressive_edu and current_section == "education":
            continue

        if current_section == "experience":
            experience_text_lines.append(ln)

    if experience_text_lines:
        print(f"üß± Collected {len(experience_text_lines)} line(s) from experience section.")
    else:
        print("‚ÑπÔ∏è No direct experience section captured; falling back to broader scan.")

    if not experience_text_lines:
        if aggressive_edu:
            filtered_lines = []
            skip = False
            for ln in lines:
                if education_keywords.search(ln):
                    skip = True
                    continue
                if re.match(r'^[A-Z ]{2,30}$', ln) and len(ln.split()) <= 4:
                    skip = False
                if not skip:
                    filtered_lines.append(ln)
            experience_text = "\n".join(filtered_lines)
        else:
            experience_text = "\n".join(lines)
    else:
        experience_text = "\n".join(experience_text_lines)

    # 3) Structured job interval pattern
    job_exp_pattern = re.compile(
        r'(?:[\w\s]+,?\s*[\w\s]*?),\s*([\w\s]+?),\s*([A-Za-z]{3}\s*\d{4})\s*[\-‚Äì‚Äîto]{1,}\s*(present|till date|current|[A-Za-z]{3}\s*\d{4})',
        re.IGNORECASE
    )
    structured_intervals = []
    for m in job_exp_pattern.finditer(text):
        sd = parse_date_any(m.group(2))
        ed_text = m.group(3)
        ed = datetime.now() if re.search(r'(present|till date|current)', ed_text, flags=re.I) else parse_date_any(ed_text)
        if sd and ed and sd < ed:
            structured_intervals.append((sd, ed))

    if structured_intervals:
        structured_intervals.sort(key=lambda x: x[0])
        merged = []
        cur_s, cur_e = structured_intervals[0]
        for s, e in structured_intervals[1:]:
            if s <= cur_e:
                cur_e = max(cur_e, e)
            else:
                merged.append((cur_s, cur_e))
                cur_s, cur_e = s, e
        merged.append((cur_s, cur_e))

        total_days = sum((e - s).days for s, e in merged)
        total_years = round(total_days / 365.25, 1)
        dbg["source"] = "structured_sections"
        dbg["structured_intervals"] = [(s.strftime("%b %Y"), e.strftime("%b %Y")) for s, e in merged]
        dbg["final_years"] = total_years
        print("üóÇÔ∏è Structured intervals (merged):")
        for s, e in merged:
            print(f" ‚Ä¢ {s.strftime('%b %Y')} ‚Üí {e.strftime('%b %Y')}")
        print(f"‚úÖ Using structured job intervals ‚Üí {total_years} years")
        EXTRACT_DEBUG_REGISTRY[fname] = dbg
        return total_years

    # 4) Generic date ranges
    intervals = []
    date_range_patterns = [
        r'([A-Za-z]{3,9}\s*\d{4})\s*[\-‚Äì‚Äîto]{1,}\s*(present|current|till date|[A-Za-z]{3,9}\s*\d{4})',
        r'(\d{4}\/\d{1,2})\s*[\-‚Äì‚Äîto]{1,}\s*(present|current|till date|\d{4}\/\d{1,2})',
        r'(\d{4})\s*[\-‚Äì‚Äîto]{1,}\s*(present|current|till date|\d{4})'
    ]
    for patt in date_range_patterns:
        for m in re.finditer(patt, experience_text, flags=re.IGNORECASE):
            sd = parse_date_any(m.group(1))
            ed = datetime.now() if re.search(r'(present|current|till date)', m.group(2), flags=re.I) else parse_date_any(m.group(2))
            if sd and ed and sd < ed:
                intervals.append((sd, ed))

    if intervals:
        intervals.sort(key=lambda x: x[0])
        merged = []
        cur_s, cur_e = intervals[0]
        for s, e in intervals[1:]:
            if s <= cur_e:
                cur_e = max(cur_e, e)
            else:
                merged.append((cur_s, cur_e))
                cur_s, cur_e = s, e
        merged.append((cur_s, cur_e))

        total_days = sum((e - s).days for s, e in merged)
        total_years = round(total_days / 365.25, 1)
        dbg["source"] = "date_ranges"
        dbg["date_range_intervals"] = [(s.strftime("%b %Y"), e.strftime("%b %Y")) for s, e in merged]
        dbg["final_years"] = total_years
        print("üóÇÔ∏è Date ranges (merged):")
        for s, e in merged:
            print(f" ‚Ä¢ {s.strftime('%b %Y')} ‚Üí {e.strftime('%b %Y')}")
        print(f"‚úÖ Using generic date ranges ‚Üí {total_years} years")
        EXTRACT_DEBUG_REGISTRY[fname] = dbg
        return total_years

    # 5) Filename hints
    if filename:
        m = re.search(r'(\d+)\s*[yY]\s*[_\-]?\s*(\d+)\s*[mM]?', filename)
        if m:
            result = round(int(m.group(1)) + int(m.group(2))/12.0, 1)
            dbg["source"] = "filename_y_m"
            dbg["final_years"] = result
            print(f"‚úÖ Using filename Y-M pattern ‚Üí {result} years (from: {filename})")
            EXTRACT_DEBUG_REGISTRY[fname] = dbg
            return result
        m = re.search(r'(\d+)\s*[yY]', filename)
        if m:
            result = float(m.group(1))
            dbg["source"] = "filename_y"
            dbg["final_years"] = result
            print(f"‚úÖ Using filename Y pattern ‚Üí {result} years (from: {filename})")
            EXTRACT_DEBUG_REGISTRY[fname] = dbg
            return result

    # 6) Simple year-year fallback
    m = re.search(r'(\d{4})\s*[\-‚Äì‚Äîto]{1,}\s*(\d{4}|present|current)', text_low)
    if m:
        sd = parse_date_any(m.group(1))
        ed = datetime.now() if re.search(r'(present|current)', m.group(2)) else parse_date_any(m.group(2))
        if sd and ed and sd < ed:
            result = round((ed - sd).days / 365.25, 1)
            dbg["source"] = "simple_year_year"
            dbg["final_years"] = result
            print(f"‚úÖ Using simple year-year match ‚Üí {result} years")
            EXTRACT_DEBUG_REGISTRY[fname] = dbg
            return result

    print("‚ùå No experience detected. Returning 0.0")
    EXTRACT_DEBUG_REGISTRY[fname] = dbg
    return 0.0

# ----------------------------------------------------------------------------------------------------------------------
# MANDATORY-FIRST SCORING helper that accepts resume_skillset and resume_exp directly
# ----------------------------------------------------------------------------------------------------------------------
def compute_score_from_sets(
    resume_skillset: set,
    resume_exp: float,
    jd_all_skills: set,
    jd_mandatory: set,
    jd_optional: set,
    jd_min_exp: int,
):
    matched_mandatory = sorted(list(jd_mandatory & resume_skillset))
    missing_mandatory = sorted(list(jd_mandatory - resume_skillset))
    matched_optional = sorted(list(jd_optional & resume_skillset))
    matched_all = sorted(list(jd_all_skills & resume_skillset))

    ratio = min(resume_exp / jd_min_exp, 1.0) if jd_min_exp > 0 else 1.0

    if jd_mandatory:
        if len(missing_mandatory) > 0:
            score = int(20 * ratio)
            return {
                "score": score,
                "matched_mandatory": matched_mandatory,
                "missing_mandatory": missing_mandatory,
                "matched_optional": matched_optional,
                "matched_all": matched_all,
                "exp_years": round(resume_exp, 1),
                "exp_met": resume_exp >= jd_min_exp if jd_min_exp > 0 else True,
                "status": "rejected",
                "reason": "Mandatory skills missing",
                "resume_skillset": sorted(list(resume_skillset)),
            }

        mand_cov = len(matched_mandatory) / max(1, len(jd_mandatory))
        opt_cov = len(matched_optional) / max(1, len(jd_optional))
        score = round(50 * mand_cov + 30 * opt_cov + 20 * ratio)
        status = "selected" if (resume_exp >= jd_min_exp if jd_min_exp > 0 else True) and score >= 50 else "rejected"
        reason = "OK" if status == "selected" else ("Experience gap" if resume_exp < jd_min_exp else "Low score")
    else:
        cov = len(matched_all) / max(1, len(jd_all_skills))
        score = round(70 * cov + 30 * ratio)
        status = "selected" if (resume_exp >= jd_min_exp if jd_min_exp > 0 else True) and score >= 50 else "rejected"
        reason = "OK" if status == "selected" else ("Experience gap" if resume_exp < jd_min_exp else "Low score")

    if resume_exp == 0.0:
        status = "needs review"
        reason = "Experience not found (parser)"

    return {
        "score": score,
        "matched_mandatory": matched_mandatory,
        "missing_mandatory": missing_mandatory,
        "matched_optional": matched_optional,
        "matched_all": matched_all,
        "exp_years": round(resume_exp, 1),
        "exp_met": resume_exp >= jd_min_exp if jd_min_exp > 0 else True,
        "status": status,
        "reason": reason,
        "resume_skillset": sorted(list(resume_skillset)),
    }

# ----------------------------------------------------------------------------------------------------------------------
# Helper: format experience string
# ----------------------------------------------------------------------------------------------------------------------
def format_exp_years(exp_float):
    years = int(exp_float)
    months = int(round((exp_float - years) * 12))
    if months == 0:
        return f"{years} years"
    return f"{years} years {months} month{'s' if months > 1 else ''}"

# ----------------------------------------------------------------------------------------------------------------------
# STREAMLIT UI 
# ----------------------------------------------------------------------------------------------------------------------


st.set_page_config(page_title="HIRE HUB ‚Äì Resume Shortlisting", page_icon="üíº", layout="wide")

APP_BG_COLOR = "#f4f8ff"

st.markdown("""
<style>
.use-jd-btn-wrap div.stButton > button:first-child {
    background-color: #16a34a !important;  /* Green */
    color: #ffffff !important;
    font-size: 18px !important;
    font-weight: bold !important;
    border-radius: 12px !important;
    padding: 12px 28px !important;
    border: none !important;
    box-shadow: 0 4px 10px rgba(0,0,0,0.2) !important;
    transition: all 0.3s ease-in-out !important;
}
.use-jd-btn-wrap div.stButton > button:first-child:hover {
    background-color: #22c55e !important;  /* Lighter green on hover */
    transform: scale(1.05);
    box-shadow: 0 6px 14px rgba(0,0,0,0.25);
}
</style>
""", unsafe_allow_html=True)

st.markdown(f"""
<style>
.stApp {{ background: linear-gradient(180deg, {APP_BG_COLOR} 0%, #ffffff 100%); }}
.header {{ background: #007acc; padding: 18px; border-radius: 12px; color: white; }}
.card {{ background: white; border-radius: 12px; padding: 12px; box-shadow: 0 2px 6px rgba(0,0,0,0.06); margin-bottom:12px; }}
.hirehub-table {{ border-collapse: collapse; width: 100%; table-layout: auto; font-family: Arial, sans-serif; font-size: 13px; }}
.hirehub-table th, .hirehub-table td {{ border: 1px solid #e2e8f0; padding: 8px; vertical-align: top; text-align: left; white-space: normal; }}
.hirehub-table th {{ background: #f1f5f9; font-weight: 600; }}
.hirehub-wrapper {{ overflow-x: auto; width: 100%; }}
.badge {{ display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 12px; font-weight: 600; color: white; }}
.badge.green {{ background-color: #16a34a; }} /* Selected */
.badge.orange {{ background-color: #f59e0b; }} /* Needs Review ‚Äî Manual */
.badge.red {{ background-color: #dc2626; }} /* Rejected */
.small-note {{ color:#64748b; font-size:12px; }}
.jd-inline-card {{
  max-width: 600px;
  margin: 10px auto 18px auto;
  background: #ffffff;
  border-radius: 18px;
  padding: 20px 18px 16px 18px;
  box-shadow: 0 12px 40px rgba(15,23,42,0.12);
}}
.header {{
  background: linear-gradient(135deg, #0ea5e9 0%, #1f6feb 60%, #9333ea 100%);
  padding: 18px 22px;
  border-radius: 12px;
  color: #fff;
  box-shadow: var(--shadow);
  position: sticky;
  top: 8px;
  z-index: 10;
}}
.jd-modal-title {{
  font-size: 20px;
  font-weight: 700;
  margin-bottom: 4px;
}}
.jd-modal-sub {{
  font-size: 13px;
  color: #6b7280;
  margin-bottom: 12px;
}}
.jd-modal-footer-text {{
  font-size: 11px;
  color: #9ca3af;
  margin-top: 4px;
}}
</style>
""", unsafe_allow_html=True)

st.markdown(f"""<div class="header"><h2 style="margin:0; font-weight:700;">üìÑ HIRE HUB ‚Äî Resume Shortlister</h2></div>""", unsafe_allow_html=True)

st.sidebar.title("‚öôÔ∏è Uploads")

# Reset & JD form state
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0
if "jd_form_payload" not in st.session_state:
    st.session_state.jd_form_payload = None
if "jd_mandatory_from_file" not in st.session_state:
    st.session_state.jd_mandatory_from_file = set()
if "show_jd_modal" not in st.session_state:
    st.session_state.show_jd_modal = False
if "results_df" not in st.session_state:
    st.session_state.results_df = None
if "results_html_table" not in st.session_state:
    st.session_state.results_html_table = None
if "show_results_now" not in st.session_state:
    st.session_state.show_results_now = False
if "last_summary" not in st.session_state:
    st.session_state.last_summary = None

if st.sidebar.button("üîÑ Reset All Uploads"):
    for key in ["jd_file", "resume_files", "resume_zip"]:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state.uploader_key += 1
    st.session_state.jd_form_payload = None
    st.session_state.jd_mandatory_from_file = set()
    st.session_state.results_df = None
    st.session_state.results_html_table = None
    st.session_state.show_results_now = False
    st.session_state.show_jd_modal = False
    st.rerun()

st.sidebar.markdown("Upload JD & multiple resumes (or a zip containing resumes).")
jd_file = st.sidebar.file_uploader(
    "üìò Upload Job Description (PDF/TXT)",
    type=['pdf', 'txt'],
    key=f"jd_{st.session_state.uploader_key}"
)
if jd_file:
    st.session_state.jd_file = jd_file

# ================== RESUME UPLOAD OPTIONS ==================
st.sidebar.subheader("üßæ Upload Resumes")
# Option 1: Multiple resume files OR folder drag-drop
resume_files = st.sidebar.file_uploader(
    "üìÇ Upload Resume Files or Drag Folder Here",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True,
    help="You can drag & drop an entire folder here (Chrome supported)"
)
if resume_files:
    st.session_state.resume_files = resume_files

# Option 2: ZIP upload (already exists)
resume_zip = st.sidebar.file_uploader(
    "üóúÔ∏è Upload Resume Folder as ZIP",
    type=["zip"]
)
if resume_zip:
    st.session_state.resume_zip = resume_zip

#aggressive_edu_exclusion = st.sidebar.checkbox("‚ö° Aggressive Education Date Exclusion", value=True)
process_button = st.sidebar.button("üöÄ Start Shortlisting", type="primary")

# Center JD Form button
st.markdown('<div style="text-align:center; margin: 18px 0 4px 0;">', unsafe_allow_html=True)
open_modal = st.button("‚úÖ Use JD Form", key="open_jd_modal", type="secondary", help="Fill JD details without a file")
st.markdown('</div>', unsafe_allow_html=True)
if open_modal:
    st.session_state.show_jd_modal = True
    st.rerun()

# JD inline form
if st.session_state.show_jd_modal:
    st.markdown('<div class="jd-modal-title">Job Description Form</div>', unsafe_allow_html=True)
    st.markdown('<div class="jd-modal-sub">Fill in JD details if you don\'t have a JD file, or want to override it.</div>', unsafe_allow_html=True)
    with st.form("jd_form_modal"):
        role_title = st.text_input("Role Title", placeholder="e.g., Embedded Linux Engineer")
        min_exp_years = st.number_input("Minimum Experience (years)", min_value=0, max_value=50, value=0, step=1)
        jd_mandatory_str = st.text_area("Mandatory Skills (comma-separated)", placeholder="e.g., Embedded C, Linux, Device Tree, U-Boot")
        jd_optional_str = st.text_area("Optional/Nice-to-have Skills (comma-separated)", placeholder="e.g., Yocto, SPI, I2C, UART")
        col_apply, col_cancel = st.columns([1,1])
        use_form = col_apply.form_submit_button("Apply JD Form", type="primary")
        cancel_form = col_cancel.form_submit_button("Cancel")

        if use_form:
            def canonize_list(s):
                items = [re.sub(r'\s+', ' ', x.strip().lower()) for x in s.split(',') if x.strip()]
                canonical = set()
                text = " " + " , ".join(items) + " "
                canonical = get_skills_from_text(text)
                for x in items:
                    if x not in canonical:
                        canonical.add(x)
                return canonical

            st.session_state.jd_form_payload = {
                "role": role_title.strip(),
                "min_exp": int(min_exp_years),
                "mandatory": canonize_list(jd_mandatory_str),
                "optional": canonize_list(jd_optional_str),
            }
            st.session_state.show_jd_modal = False
            st.toast("JD Form applied.", icon="‚úÖ")
            st.rerun()

        if cancel_form:
            st.session_state.show_jd_modal = False
            st.rerun()

    st.markdown('<div class="jd-modal-footer-text">Tip: Once applied, this JD form will override any uploaded JD file unless you reset uploads.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# Tabs
tab1, tab2, tab3 = st.tabs(["üìä Results", "üìò Job Description", "üßæ Resumes"])

with tab2:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 1 ‚Äî Job Description")

    jd_text = ""
    jd_from_form = st.session_state.jd_form_payload is not None
    if jd_from_form:
        p = st.session_state.jd_form_payload
        st.success("Using JD **Form** (overrides file).")
        st.markdown(f"**Role**: {p['role'] or '‚Äî'}")
        st.markdown(f"**Min Experience**: **{p['min_exp']} years**")
        st.markdown(f"**Mandatory (form)**: {', '.join(sorted(list(p['mandatory']))) or '‚Äî'}")
        st.markdown(f"**Optional (form)**: {', '.join(sorted(list(p['optional']))) or '‚Äî'}")
        jd_keywords = (p["mandatory"] | p["optional"])
        jd_min_exp = p["min_exp"]
        st.info("You can still upload a JD file, but the form values from the center-button JD Form will be used.")
    else:
        if jd_file:
            if jd_file.type == "application/pdf":
                jd_text = extract_text_from_pdf(jd_file)
            else:
                jd_text = extract_text_from_txt(jd_file)
            st.text_area("JD Preview (first 2000 chars):", jd_text[:2000], height=220)

            # Extract skills from file (local canonical extractor)
            jd_keywords = get_keywords_from_jd(jd_text)

            # Try to infer min exp from text
            jd_min_exp = 0
            for pat in [r'(\d+)\s*\+\s*years', r'(\d+)\s*\-\s*\d+\s*years',
                        r'minimum\s*of\s*(\d+)\s*years', r'at\s*least\s*(\d+)\s*years',
                        r'(\d+)\s*years\s*of\s*experience']:
                found = re.findall(pat, jd_text)
                if found:
                    try:
                        jd_min_exp = int(found[0])
                        break
                    except:
                        pass

            st.markdown("#### Pick Mandatory Skills (from detected)")
            picked = st.multiselect(
                "Mark mandatory skills (optional ‚Äî used as gate)",
                sorted(list(jd_keywords)),
                default=sorted(list(st.session_state.jd_mandatory_from_file or []))
            )
            st.caption("These will be checked first. If any are missing in a resume ‚Üí Rejected.")
            if st.button("Apply Mandatory Selection"):
                st.session_state.jd_mandatory_from_file = set(picked)
                st.success("Saved mandatory selection.")

            st.markdown(f"**Detected JD Skills** ({len(jd_keywords)}): " + (", ".join(sorted(list(jd_keywords))) if jd_keywords else "‚Äî"))
            st.markdown(f"**Min Experience (parsed)**: **{jd_min_exp} years**")
        else:
            jd_keywords = set()
            jd_min_exp = 0
            st.info("Upload JD (PDF/TXT) or click **Use JD Form** in the center of the page.")

    st.markdown('</div>', unsafe_allow_html=True)

with tab3:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Step 2 ‚Äî Upload Resumes")
    if st.session_state.get("resume_files"):
        st.success(f"{len(st.session_state.resume_files)} resume(s) uploaded (individual).")
        for r in st.session_state.resume_files:
            st.markdown(f"- üìÑ **{r.name}**")
    if st.session_state.get("resume_zip"):
        st.success("Zip uploaded.")
        st.markdown(f"- üì¶ **{st.session_state.resume_zip.name}**")
    if not st.session_state.get("resume_files") and not st.session_state.get("resume_zip"):
        st.info("No resumes uploaded yet.")
    st.markdown('</div>', unsafe_allow_html=True)

with tab1:
# Result renderer (unchanged)
    def render_results_block(df, jd_min_exp_value, jd_all_skills, jd_mandatory, jd_optional, key_suffix: str = "default"):
        st.success(f"Processed {len(df)} resume(s).")
        if not df.empty:
            top_k = df.head(3)
            cols = st.columns(3)
            for c, (_, r) in zip(cols, top_k.iterrows()):
                c.markdown(f"**{r['Candidate Name']}**", unsafe_allow_html=True)
                c.markdown(f"Score: **{r['Score']}** ‚Ä¢ Exp: **{r['Years Experience']}**", unsafe_allow_html=True)
                c.markdown(r['Status'], unsafe_allow_html=True)
                c.caption((r['Matched Mandatory'][:120] + ("..." if len(r['Matched Mandatory']) > 120 else "")) or "‚Äî")

            st.markdown("### Shortlisted Candidates")

            html_table = df[[
                "Candidate Name",
                "Score",
                "Years Experience",
                "Status",
                "Matched Mandatory Count",
                "Mandatory Missing Count",
                "Matched Mandatory",
                "Mandatory Missing",
                "Matched Optional",
                "Matched (All JD Skills)",
                "Unmatched (All JD Skills)",
            ]].to_html(index=False, escape=False, classes="hirehub-table")
            wrapper = f'<div class="hirehub-wrapper">{html_table}</div>'
            st.markdown(wrapper, unsafe_allow_html=True)

            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Shortlisted')
            buf.seek(0)

            st.download_button(
                "üì• Download Shortlisted Candidates (Excel)",
                buf,
                file_name="shortlisted_candidates.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"download_excel_{key_suffix}",
            )
            st.session_state.results_html_table = wrapper
        else:
            st.info("No candidates to display after filters.")

    # Helper: read zip
    def read_zip_to_files(zip_file):
        try:
            zip_file.seek(0)
            z = zipfile.ZipFile(io.BytesIO(zip_file.read()))
            files = []
            for name in z.namelist():
                if name.lower().endswith(('.pdf', '.docx', '.txt')):
                    try:
                        data = z.read(name)
                        bio = io.BytesIO(data)
                        bio.name = os.path.basename(name)
                        files.append(bio)
                    except:
                        continue
            return files
        except Exception as e:
            st.error(f"Failed to process zip file: {e}")
            return []

    # ----------------------------------------------------------------------------------------------------------------------
    # PROCESS ACTION ‚Äî now using LLM-first flow with fallbacks to local extraction
    # ----------------------------------------------------------------------------------------------------------------------
    if process_button:
        combined_resumes = list(st.session_state.get("resume_files") or [])
        if st.session_state.get("resume_zip"):
            combined_resumes.extend(read_zip_to_files(st.session_state.resume_zip))

        jd_from_form = st.session_state.jd_form_payload is not None

        if not (jd_from_form or st.session_state.get("jd_file")):
            st.warning("Please provide a Job Description via Form or File.")
        elif not combined_resumes:
            st.warning("Please upload at least one resume (individual files or a zip).")
        else:
            # Resolve JD payload (same as before)
            if jd_from_form:
                p = st.session_state.jd_form_payload
                jd_min_exp = p["min_exp"]
                jd_mandatory = set(p["mandatory"])
                jd_optional = set(p["optional"])
                jd_all_skills = jd_mandatory | jd_optional
            else:
                jd_file = st.session_state.jd_file
                if jd_file.type == "application/pdf":
                    jd_text = extract_text_from_pdf(jd_file)
                else:
                    jd_text = extract_text_from_txt(jd_file)
                jd_all_skills = get_keywords_from_jd(jd_text)
                jd_mandatory = set(st.session_state.jd_mandatory_from_file or set())
                jd_optional = jd_all_skills - jd_mandatory
                jd_min_exp = 0
                for pat in [r'(\d+)\s*\+\s*years', r'(\d+)\s*\-\s*\d+\s*years',
                            r'minimum\s*of\s*(\d+)\s*years', r'at\s*least\s*(\d+)\s*years',
                            r'(\d+)\s*years\s*of\s*experience']:
                    found = re.findall(pat, jd_text)
                    if found:
                        try:
                            jd_min_exp = int(found[0])
                            break
                        except:
                            pass

            # ------------------ Process resumes (with percent indicator) ------------------
            
            results = []
            total = len(combined_resumes)

            # progress UI: progress bar + percentage label
            progress_bar = st.progress(0)
            progress_label = st.empty()  # we'll use this to show "xx% ‚Äî n/total processed"
            def _update_progress(i, total):
                if total <= 0:
                    pct = 100
                else:
                    pct = int(((i) / total) * 100)
                    if pct > 100:
                        pct = 100
                # update progress bar (expects fraction 0..1)
                frac = min(max(i / max(1, total), 0.0), 1.0)
                progress_bar.progress(frac)
                # show percentage with processed count
                progress_label.markdown(f"**Progress:** {pct}% ‚Äî {i}/{total} processed")

            # initialize display
            _update_progress(0, total)

            def status_badge(status, reason=""):
                color = {"selected": "green", "needs review": "orange", "rejected": "red"}[status]
                label = {
                    "selected": "Selected",
                    "rejected": "Rejected",
                    "needs review": "Needs Review ‚Äî Manual Check"
                }[status]
                return f'<span class="badge {color}">{label}</span>' + (f' <span class="small-note">({reason})</span>' if reason else "")

            for idx, rf in enumerate(combined_resumes, start=1):
                fname = getattr(rf, "name", f"resume_{idx}")
                txt = ""
                try:
                    if fname.lower().endswith(".pdf"):
                        txt = extract_text_from_pdf(rf)
                    elif fname.lower().endswith(".docx"):
                        txt = extract_text_from_docx(rf)
                    else:
                        txt = extract_text_from_txt(rf)
                except Exception:
                    txt = ""

                # 1) Skills extraction: ask LLM first, fallback to local get_skills_from_text
                skills_list, raw_sk = call_gemini_for_skills(txt)
                if not skills_list:
                    # fallback to local extractor
                    local_skills = get_skills_from_text(txt)
                    skills_set = set(local_skills)
                else:
                    skills_set = set(skills_list)

                # 2) Experience extraction: ask LLM first, fallback to local extractor
                years_decimal, raw_exp = call_gemini_for_years(txt)
                if years_decimal == 0.0:
                    # local fallback
                    years_decimal = extract_experience_from_resume(txt, filename=fname)

                # 3) Compute score using sets (ensures consistent response shape)
                scoring = compute_score_from_sets(
                    resume_skillset=skills_set,
                    resume_exp=years_decimal,
                    jd_all_skills=jd_all_skills,
                    jd_mandatory=jd_mandatory,
                    jd_optional=jd_optional,
                    jd_min_exp=jd_min_exp
                )

                # (your debug / print block left unchanged)
                src_info = EXTRACT_DEBUG_REGISTRY.get(fname, {})
                src = src_info.get("source", "n/a")
                print("\n" + "-"*78)
                print(f"üìÑ RESUME REPORT: {fname}")
                print("-"*78)
                print(f"‚Ä¢ Matched Mandatory ({len(scoring['matched_mandatory'])}): {', '.join(scoring['matched_mandatory']) or '‚Äî'}")
                print(f"‚Ä¢ Missing Mandatory ({len(scoring['missing_mandatory'])}): {', '.join(scoring['missing_mandatory']) or '‚Äî'}")
                if jd_optional:
                    print(f"‚Ä¢ Matched Optional ({len(scoring['matched_optional'])}): {', '.join(scoring['matched_optional']) or '‚Äî'}")
                print(f"‚Ä¢ Experience : {format_exp_years(scoring['exp_years'])} (source: {src})")
                if src_info.get("structured_intervals"):
                    print("‚Ä¢ Intervals (structured):")
                    for s,e in src_info["structured_intervals"]:
                        print(f" - {s} ‚Üí {e}")
                if src_info.get("date_range_intervals"):
                    print("‚Ä¢ Intervals (date ranges):")
                    for s,e in src_info["date_range_intervals"]:
                        print(f" - {s} ‚Üí {e}")
                if src_info.get("explicit_values"):
                    print(f"‚Ä¢ Explicit numeric mentions: {src_info['explicit_values']}")
                print(f"‚Ä¢ JD Min Exp : {jd_min_exp} years")
                print(f"‚Ä¢ Score : {scoring['score']}")
                print(f"‚Ä¢ Status : {scoring['status'].upper()} ({scoring['reason']})")
                print("-"*78)

                missing_all = sorted(list((jd_all_skills - set(scoring["matched_all"])))) if jd_all_skills else []

                results.append({
                    "Candidate Name": os.path.splitext(fname)[0],
                    "Score": scoring["score"],
                    "Years Experience": format_exp_years(scoring["exp_years"]),
                    "Status": status_badge(scoring["status"], scoring["reason"]),
                    "Matched Mandatory Count": len(scoring["matched_mandatory"]),
                    "Mandatory Missing Count": len(scoring["missing_mandatory"]),
                    "Matched Mandatory": ", ".join(scoring["matched_mandatory"]),
                    "Mandatory Missing": ", ".join(scoring["missing_mandatory"]),
                    "Matched Optional": ", ".join(scoring["matched_optional"]),
                    "Matched (All JD Skills)": ", ".join(scoring["matched_all"]),
                    "Unmatched (All JD Skills)": ", ".join(missing_all),
                    "Filename": fname
                })

                # update progress display (pass idx so we show processed count)
                _update_progress(idx, total)

            # finalize progress UI
            _update_progress(total, total)
            progress_label.markdown(f"**Progress:** 100% ‚Äî Completed ({total}/{total})")

            df = pd.DataFrame(sorted(results, key=lambda x: (x['Score'], -x["Mandatory Missing Count"]), reverse=True))
            st.session_state.results_df = df

            if not df.empty:
                tk = df.head(3)
                st.session_state.last_summary = tk.to_dict('records')
            else:
                st.session_state.last_summary = None

            
            st.session_state.show_results_now = True
            st.toast("Shortlisting complete. Showing results‚Ä¶", icon="üìä")
            st.rerun()
            

    # Auto show results area
    if st.session_state.show_results_now and st.session_state.results_df is not None:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("Results")
        render_results_block(
            st.session_state.results_df,
            0, set(), set(), set(),
            key_suffix="autoshow"
        )
        st.markdown('</div>', unsafe_allow_html=True)